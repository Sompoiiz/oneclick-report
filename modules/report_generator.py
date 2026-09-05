"""
โมดูลสร้างรายงานสรุปผลการประเมินโครงการเป็นไฟล์ Word (.docx)
โดยใช้ python-docx และคืนค่าเป็น BytesIO เพื่อให้ Streamlit ดาวน์โหลดได้ทันที
โดยไม่ต้องเขียนไฟล์ลงดิสก์

โครงสร้างรายงานอิงตามรูปแบบเอกสารสรุปผลการประเมินโครงการที่ใช้งานจริง
(ปก + 1.รูปแบบการประเมินผล + 2.เกณฑ์การประเมิน + 3.ผลการประเมิน (ข้อมูลผู้ตอบ +
ผลรายด้าน) + 4.ข้อเสนอแนะอื่นๆ)
"""
from __future__ import annotations

import io

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from .stats_utils import DEFAULT_GROUP_NAME, INTERPRETATION_LEVELS

# ฟอนต์มาตรฐานเอกสารราชการ/วิชาการไทย (ตรงกับที่ใช้ในเอกสารตัวอย่างจริง)
THAI_FONT = "TH SarabunPSK"


def _set_font(run, size: int = 16, bold: bool = False) -> None:
    """
    ตั้งค่าฟอนต์ไทยให้ครบทั้ง ascii/eastAsia/cs
    (จุดที่มักพลาด: ถ้าตั้งแค่ run.font.name เฉย ๆ Word บางเครื่องจะ fallback
    ไปใช้ฟอนต์อื่นกับตัวอักษรไทย ทำให้สระ/วรรณยุกต์เพี้ยน จึงต้องตั้ง rFonts ทุก tag)
    """
    run.font.name = THAI_FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), THAI_FONT)
    rfonts.set(qn("w:cs"), THAI_FONT)


def _add_paragraph(doc: Document, text: str, size: int = 16, bold: bool = False, center: bool = False):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    _set_font(run, size=size, bold=bold)
    return p


def _style_cell(cell, text, size: int = 14, bold: bool = False, center: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(text))
    _set_font(run, size=size, bold=bold)


def _add_table_header(table, headers: list[str]) -> None:
    for i, h in enumerate(headers):
        _style_cell(table.rows[0].cells[i], h, bold=True, center=True)


def generate_report(
    project_name: str,
    event_period: str,
    event_location: str,
    respondent_count: int,
    demographics: dict[str, pd.DataFrame],
    groups: list[dict],
    feedback_texts: list[str],
) -> io.BytesIO:
    """สร้างเอกสาร Word สรุปผลการประเมิน คืนค่าเป็น BytesIO พร้อมดาวน์โหลด"""
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    is_single_group = len(groups) == 1 and groups[0]["group_name"] == DEFAULT_GROUP_NAME

    # ---------- ปก ----------
    _add_paragraph(doc, "สรุปผลการประเมินผล", size=18, bold=True, center=True)
    if project_name and project_name != "-":
        _add_paragraph(doc, project_name, size=18, bold=True, center=True)
    if event_period:
        _add_paragraph(doc, f"ระหว่างวันที่ {event_period}", size=18, bold=True, center=True)
    if event_location:
        _add_paragraph(doc, f"ณ {event_location}", size=18, bold=True, center=True)
    _add_paragraph(doc, "…" * 22, size=18, bold=True, center=True)
    doc.add_paragraph()

    # ---------- 1. รูปแบบการประเมินผล ----------
    _add_paragraph(doc, "1. รูปแบบการประเมินผล", size=16, bold=True)

    toc_parts = []
    if demographics:
        toc_parts.append("ข้อมูลของผู้ตอบแบบประเมิน")
    if is_single_group:
        toc_parts.append("ผลการประเมิน")
    else:
        toc_parts.extend(g["group_name"] for g in groups)
    if feedback_texts is not None:
        toc_parts.append("ข้อเสนอแนะอื่นๆ")

    _add_paragraph(
        doc,
        f"เครื่องมือที่ใช้ประเมิน ได้แก่ แบบประเมิน 1 ชุด มี {len(toc_parts)} ตอน "
        "(แบบออนไลน์ Google Form)",
    )
    for i, part in enumerate(toc_parts, start=1):
        _add_paragraph(doc, f"ตอนที่ {i} {part}")
    _add_paragraph(doc, f"จำนวนผู้ตอบแบบประเมินทั้งหมด {respondent_count} คน")
    doc.add_paragraph()

    # ---------- 2. เกณฑ์ในการประเมินกิจกรรม ----------
    _add_paragraph(doc, "2. เกณฑ์ในการประเมินกิจกรรม", size=16, bold=True)
    _add_paragraph(doc, "ในการประเมินครั้งนี้กำหนดเกณฑ์การแปลความหมายค่าเฉลี่ยไว้ดังนี้")
    crit_table = doc.add_table(rows=1, cols=2)
    crit_table.style = "Table Grid"
    _add_table_header(crit_table, ["ค่าเฉลี่ย", "ระดับความคิดเห็น/ความพึงพอใจ"])
    for low, high, label in INTERPRETATION_LEVELS:
        cells = crit_table.add_row().cells
        _style_cell(cells[0], f"{low:.2f}-{high:.2f}", center=True)
        _style_cell(cells[1], label, center=True)
    doc.add_paragraph()

    # ---------- 3. ผลการประเมินกิจกรรม ----------
    _add_paragraph(doc, "3. ผลการประเมินกิจกรรม", size=16, bold=True)

    sub_no = 1

    # 3.x ข้อมูลของผู้ตอบแบบประเมิน
    if demographics:
        _add_paragraph(doc, f"3.{sub_no} ข้อมูลของผู้ตอบแบบประเมิน", size=16, bold=True)
        sub_no += 1
        for col_name, breakdown_df in demographics.items():
            _add_paragraph(doc, f"ตารางแสดงจำนวนและร้อยละของผู้ตอบแบบประเมิน จำแนกตาม{col_name}", bold=True)
            demo_table = doc.add_table(rows=1, cols=len(breakdown_df.columns))
            demo_table.style = "Table Grid"
            _add_table_header(demo_table, list(breakdown_df.columns))
            for _, row in breakdown_df.iterrows():
                cells = demo_table.add_row().cells
                for i, col in enumerate(breakdown_df.columns):
                    _style_cell(cells[i], row[col], center=(i != 0))
            doc.add_paragraph()

    # 3.x ผลการประเมินรายด้าน
    if groups:
        _add_paragraph(doc, f"3.{sub_no} ผลการประเมินรายด้าน", size=16, bold=True)

        if is_single_group:
            g = groups[0]
            table = doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            _add_table_header(table, ["รายการ", "ค่าเฉลี่ย", "S.D.", "การแปลความหมาย"])
            for _, item in g["items"].iterrows():
                cells = table.add_row().cells
                _style_cell(cells[0], item["ข้อคำถาม"])
                _style_cell(cells[1], f"{item['ค่าเฉลี่ย (Mean)']:.2f}", center=True)
                _style_cell(cells[2], f"{item['ส่วนเบี่ยงเบนมาตรฐาน (S.D.)']:.2f}", center=True)
                _style_cell(cells[3], item["แปลผล"], center=True)
            avg = g["average"]
            avg_cells = table.add_row().cells
            _style_cell(avg_cells[0], "ค่าเฉลี่ยรวม", bold=True)
            _style_cell(avg_cells[1], f"{avg['mean']:.2f}", bold=True, center=True)
            _style_cell(avg_cells[2], f"{avg['sd']:.2f}", bold=True, center=True)
            _style_cell(avg_cells[3], avg["level"], bold=True, center=True)
            doc.add_paragraph()
        else:
            # ตารางเดียวต่อเนื่องกันทุกกลุ่ม (ตรงกับรูปแบบเอกสารตัวอย่างจริง) แทนที่จะแยกตารางต่อกลุ่ม
            table = doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            _add_table_header(table, ["รายการ", "ค่าเฉลี่ย", "S.D.", "การแปลความหมาย"])

            for gi, g in enumerate(groups, start=1):
                header_cells = table.add_row().cells
                merged = header_cells[0].merge(header_cells[1]).merge(header_cells[2]).merge(header_cells[3])
                _style_cell(merged, f"{gi}. {g['group_name']}", bold=True)

                for _, item in g["items"].iterrows():
                    cells = table.add_row().cells
                    _style_cell(cells[0], item["ข้อคำถาม"])
                    _style_cell(cells[1], f"{item['ค่าเฉลี่ย (Mean)']:.2f}", center=True)
                    _style_cell(cells[2], f"{item['ส่วนเบี่ยงเบนมาตรฐาน (S.D.)']:.2f}", center=True)
                    _style_cell(cells[3], item["แปลผล"], center=True)

                avg = g["average"]
                avg_cells = table.add_row().cells
                _style_cell(avg_cells[0], f"ค่าเฉลี่ยด้าน{g['group_name']}", bold=True)
                _style_cell(avg_cells[1], f"{avg['mean']:.2f}", bold=True, center=True)
                _style_cell(avg_cells[2], f"{avg['sd']:.2f}", bold=True, center=True)
                _style_cell(avg_cells[3], avg["level"], bold=True, center=True)

            doc.add_paragraph()

    # ---------- 4. ข้อเสนอแนะอื่นๆ ----------
    if feedback_texts:
        _add_paragraph(doc, "4. ข้อเสนอแนะอื่น ๆ", size=16, bold=True)
        _add_paragraph(doc, "ผู้เข้าร่วมกิจกรรม/โครงการได้เสนอข้อคิดเห็นอื่น ๆ ไว้ดังนี้")
        for text in feedback_texts:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(text)
            _set_font(run, size=16)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
